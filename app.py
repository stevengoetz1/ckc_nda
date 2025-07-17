from flask import Flask, render_template, request, send_file
from docx import Document
import os
import io

app = Flask(__name__)

def replace_text_across_runs(paragraph, search, replace):
    # Concatenate all run texts
    full_text = ''.join(run.text for run in paragraph.runs)
    idx = full_text.find(search)
    if idx == -1:
        return False
    # Find which runs the search string spans
    run_indices = []
    char_count = 0
    for i, run in enumerate(paragraph.runs):
        next_count = char_count + len(run.text)
        if char_count <= idx < next_count or char_count < idx + len(search) <= next_count or (idx < char_count and next_count <= idx + len(search)):
            run_indices.append(i)
        char_count = next_count
    # Remove the runs that contain the search string
    first = run_indices[0]
    last = run_indices[-1]
    before = ''.join(run.text for run in paragraph.runs[:first])
    after = ''.join(run.text for run in paragraph.runs[last+1:])
    # Save formatting from the first run
    style = paragraph.runs[first].style
    font = paragraph.runs[first].font
    # Remove all runs in the range
    for _ in range(last - first + 1):
        paragraph._element.remove(paragraph.runs[first]._element)
    # Insert new run with replacement
    new_run = paragraph.add_run(replace)
    new_run.style = style
    new_run.font.bold = font.bold
    new_run.font.italic = font.italic
    new_run.font.underline = font.underline
    new_run.font.size = font.size
    new_run.font.name = font.name
    # Rebuild the paragraph text
    paragraph.text = before + full_text[:idx] + replace + full_text[idx+len(search):] + after
    return True

def replace_in_paragraph(paragraph, replacements):
    for search, replace in replacements:
        while True:
            changed = replace_text_across_runs(paragraph, search, replace)
            if not changed:
                break

def replace_in_table(table, replacements):
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                replace_in_paragraph(para, replacements)

def replace_text_preserving_formatting(doc, full_name, firm_name, project_name):
    first_name = full_name.strip().split()[0]
    replacements = [
        ("Finley Bond", full_name),
        ("Burlington Street Partners", firm_name),
        ("Project Slab", f"Project {project_name}"),
        ("Slab", project_name),
        ("Dear Finley,", f"Dear {first_name},")
    ]
    for para in doc.paragraphs:
        replace_in_paragraph(para, replacements)
    for table in doc.tables:
        replace_in_table(table, replacements)
    return doc

@app.route("/", methods=["GET", "POST"])
def form():
    if request.method == "POST":
        full_name = request.form["full_name"]
        firm_name = request.form["firm_name"]
        project_name = request.form["project_name"]

        doc_path = os.path.join(os.path.dirname(__file__), "Project Slab_NDA_Burlington Street Partners.docx")
        doc = Document(doc_path)
        updated_doc = replace_text_preserving_formatting(doc, full_name, firm_name, project_name)

        output_stream = io.BytesIO()
        updated_doc.save(output_stream)
        output_stream.seek(0)
        return send_file(output_stream, as_attachment=True, download_name=f"NDA_{firm_name.replace(' ', '_')}.docx")

    return render_template("form.html")

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)
